"""Document processing and ingestion pipeline."""

import asyncio
import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import uuid4

import aiofiles
from docx import Document as DocxDocument
from pypdf2 import PdfReader

from ..chunking.strategies import chunk_document
from ..config.settings import get_settings
from ..models.base import (
    ChunkingConfig,
    Document,
    DocumentMetadata,
    DocumentType,
    IngestionJob,
)
from ..utils.exceptions import DocumentProcessingError, IngestionError
from ..utils.helpers import (
    clean_text,
    extract_text_from_html,
    get_file_hash,
    get_file_size,
    get_file_type,
    parse_metadata_from_filename,
    sanitize_filename,
    validate_file_type,
)
from ..utils.logging import log_ingestion, RAGLogger


class DocumentProcessor:
    """Base document processor class."""
    
    def __init__(self):
        self.settings = get_settings()
    
    async def process_file(
        self, 
        file_path: str,
        metadata: Optional[DocumentMetadata] = None
    ) -> Tuple[Document, List[Any]]:
        """Process a single file and return document with chunks."""
        
        # Validate file
        if not os.path.exists(file_path):
            raise DocumentProcessingError(f"File not found: {file_path}")
        
        if not validate_file_type(file_path, self.settings.ingestion.supported_extensions):
            raise DocumentProcessingError(f"Unsupported file type: {file_path}")
        
        file_size = get_file_size(file_path)
        if file_size > self.settings.ingestion.max_file_size:
            raise DocumentProcessingError(
                f"File too large: {file_size} bytes (max: {self.settings.ingestion.max_file_size})"
            )
        
        # Determine document type
        doc_type = get_file_type(file_path)
        if not doc_type:
            raise DocumentProcessingError(f"Could not determine document type: {file_path}")
        
        # Extract content based on file type
        with RAGLogger("document_extraction", {"file_path": file_path, "type": doc_type}):
            content = await self._extract_content(file_path, doc_type)
        
        # Parse metadata from filename if not provided
        if metadata is None:
            filename_metadata = parse_metadata_from_filename(os.path.basename(file_path))
            metadata = DocumentMetadata(
                app=filename_metadata.get("app", "unknown"),
                version=filename_metadata.get("version", "1.0"),
                audience=filename_metadata.get("audience", "internal"),
                last_reviewed=datetime.utcnow(),
                tags=[],
                source_url=None,
                department=None,
                sensitivity="internal"
            )
        
        # Create document
        document = Document(
            title=metadata.title if hasattr(metadata, 'title') else Path(file_path).stem,
            content=content,
            doc_type=doc_type,
            metadata=metadata,
            file_path=file_path,
            file_size=file_size,
            checksum=get_file_hash(file_path)
        )
        
        # Generate chunks
        with RAGLogger("document_chunking", {"document_id": str(document.id)}):
            chunks = chunk_document(document, self.settings.chunking)
        
        return document, chunks
    
    async def _extract_content(self, file_path: str, doc_type: DocumentType) -> str:
        """Extract text content from file based on type."""
        try:
            if doc_type == DocumentType.PDF:
                return await self._extract_pdf_content(file_path)
            elif doc_type == DocumentType.DOCX:
                return await self._extract_docx_content(file_path)
            elif doc_type == DocumentType.TXT:
                return await self._extract_text_content(file_path)
            elif doc_type == DocumentType.MD:
                return await self._extract_markdown_content(file_path)
            elif doc_type == DocumentType.HTML:
                return await self._extract_html_content(file_path)
            elif doc_type == DocumentType.JSON:
                return await self._extract_json_content(file_path)
            else:
                raise DocumentProcessingError(f"Unsupported document type: {doc_type}")
                
        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract content from {file_path}: {str(e)}")
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF file."""
        loop = asyncio.get_event_loop()
        
        def _read_pdf():
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text_parts = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return "\n\n".join(text_parts)
        
        content = await loop.run_in_executor(None, _read_pdf)
        return clean_text(content)
    
    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        loop = asyncio.get_event_loop()
        
        def _read_docx():
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            return "\n\n".join(text_parts)
        
        content = await loop.run_in_executor(None, _read_docx)
        return clean_text(content)
    
    async def _extract_text_content(self, file_path: str) -> str:
        """Extract content from plain text file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        return clean_text(content)
    
    async def _extract_markdown_content(self, file_path: str) -> str:
        """Extract content from Markdown file."""
        # For now, treat as plain text - could be enhanced with markdown parsing
        return await self._extract_text_content(file_path)
    
    async def _extract_html_content(self, file_path: str) -> str:
        """Extract text content from HTML file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            html_content = await file.read()
        
        text_content = extract_text_from_html(html_content)
        return clean_text(text_content)
    
    async def _extract_json_content(self, file_path: str) -> str:
        """Extract content from JSON file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        
        try:
            # Parse JSON and convert to readable text
            data = json.loads(content)
            
            def json_to_text(obj, level=0):
                """Convert JSON object to readable text."""
                indent = "  " * level
                lines = []
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            lines.append(f"{indent}{key}:")
                            lines.append(json_to_text(value, level + 1))
                        else:
                            lines.append(f"{indent}{key}: {value}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        if isinstance(item, (dict, list)):
                            lines.append(f"{indent}[{i}]:")
                            lines.append(json_to_text(item, level + 1))
                        else:
                            lines.append(f"{indent}[{i}]: {item}")
                else:
                    lines.append(f"{indent}{obj}")
                
                return "\n".join(lines)
            
            readable_text = json_to_text(data)
            return clean_text(readable_text)
            
        except json.JSONDecodeError:
            # If not valid JSON, treat as plain text
            return clean_text(content)


class BatchIngestionService:
    """Service for batch processing multiple documents."""
    
    def __init__(self):
        self.settings = get_settings()
        self.processor = DocumentProcessor()
    
    async def process_directory(
        self,
        directory_path: str,
        metadata_template: Optional[DocumentMetadata] = None,
        recursive: bool = True
    ) -> IngestionJob:
        """Process all supported files in a directory."""
        
        job = IngestionJob(
            file_path=directory_path,
            status="running"
        )
        
        try:
            # Find all supported files
            files = self._find_supported_files(directory_path, recursive)
            
            if not files:
                job.status = "completed"
                job.completed_at = datetime.utcnow()
                return job
            
            # Process files in batches
            processed_docs = 0
            total_chunks = 0
            
            for batch in self._batch_files(files, self.settings.ingestion.batch_size):
                batch_results = await self._process_file_batch(batch, metadata_template)
                
                for document, chunks in batch_results:
                    processed_docs += 1
                    total_chunks += len(chunks)
                    
                    # Here you would typically store document and chunks in database
                    # and vector store - this will be implemented in the storage layer
            
            job.document_count = processed_docs
            job.chunk_count = total_chunks
            job.status = "completed"
            job.completed_at = datetime.utcnow()
            
            log_ingestion(
                file_path=directory_path,
                status="success",
                document_count=processed_docs,
                chunk_count=total_chunks,
                duration=(job.completed_at - job.started_at).total_seconds()
            )
            
        except Exception as e:
            job.status = "failed"
            job.error_message = str(e)
            job.completed_at = datetime.utcnow()
            
            log_ingestion(
                file_path=directory_path,
                status="error",
                error=str(e)
            )
            
            raise IngestionError(f"Batch ingestion failed: {str(e)}")
        
        return job
    
    async def process_files(
        self,
        file_paths: List[str],
        metadata_template: Optional[DocumentMetadata] = None
    ) -> IngestionJob:
        """Process a list of specific files."""
        
        job = IngestionJob(
            file_path=",".join(file_paths),
            status="running"
        )
        
        try:
            processed_docs = 0
            total_chunks = 0
            
            # Process files in batches
            for batch in self._batch_files(file_paths, self.settings.ingestion.batch_size):
                batch_results = await self._process_file_batch(batch, metadata_template)
                
                for document, chunks in batch_results:
                    processed_docs += 1
                    total_chunks += len(chunks)
            
            job.document_count = processed_docs
            job.chunk_count = total_chunks
            job.status = "completed"
            job.completed_at = datetime.utcnow()
            
        except Exception as e:
            job.status = "failed"
            job.error_message = str(e)
            job.completed_at = datetime.utcnow()
            raise IngestionError(f"File ingestion failed: {str(e)}")
        
        return job
    
    def _find_supported_files(self, directory_path: str, recursive: bool = True) -> List[str]:
        """Find all supported files in directory."""
        supported_extensions = self.settings.ingestion.supported_extensions
        files = []
        
        if recursive:
            for root, _, filenames in os.walk(directory_path):
                for filename in filenames:
                    if any(filename.lower().endswith(ext) for ext in supported_extensions):
                        files.append(os.path.join(root, filename))
        else:
            for filename in os.listdir(directory_path):
                file_path = os.path.join(directory_path, filename)
                if os.path.isfile(file_path):
                    if any(filename.lower().endswith(ext) for ext in supported_extensions):
                        files.append(file_path)
        
        return files
    
    def _batch_files(self, files: List[str], batch_size: int) -> List[List[str]]:
        """Split files into batches."""
        return [files[i:i + batch_size] for i in range(0, len(files), batch_size)]
    
    async def _process_file_batch(
        self,
        file_paths: List[str],
        metadata_template: Optional[DocumentMetadata] = None
    ) -> List[Tuple[Document, List[Any]]]:
        """Process a batch of files concurrently."""
        
        # Create semaphore to limit concurrent processing
        semaphore = asyncio.Semaphore(self.settings.ingestion.concurrent_workers)
        
        async def process_single_file(file_path: str):
            async with semaphore:
                # Create metadata for this file
                file_metadata = metadata_template
                if file_metadata is None:
                    filename_metadata = parse_metadata_from_filename(os.path.basename(file_path))
                    file_metadata = DocumentMetadata(
                        app=filename_metadata.get("app", "unknown"),
                        version=filename_metadata.get("version", "1.0"),
                        audience=filename_metadata.get("audience", "internal"),
                        last_reviewed=datetime.utcnow(),
                        tags=[],
                        source_url=None,
                        department=None,
                        sensitivity="internal"
                    )
                
                return await self.processor.process_file(file_path, file_metadata)
        
        # Process all files concurrently
        tasks = [process_single_file(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out exceptions and return successful results
        successful_results = []
        for result in results:
            if isinstance(result, Exception):
                # Log the error but continue processing other files
                log_ingestion(
                    file_path="batch_file",
                    status="error",
                    error=str(result)
                )
            else:
                successful_results.append(result)
        
        return successful_results


class IngestionService:
    """Main ingestion service orchestrating document processing."""
    
    def __init__(self):
        self.batch_service = BatchIngestionService()
        self.processor = DocumentProcessor()
    
    async def ingest_file(
        self,
        file_path: str,
        metadata: Optional[DocumentMetadata] = None
    ) -> Tuple[Document, List[Any]]:
        """Ingest a single file."""
        return await self.processor.process_file(file_path, metadata)
    
    async def ingest_directory(
        self,
        directory_path: str,
        metadata_template: Optional[DocumentMetadata] = None,
        recursive: bool = True
    ) -> IngestionJob:
        """Ingest all files in a directory."""
        return await self.batch_service.process_directory(
            directory_path, metadata_template, recursive
        )
    
    async def ingest_files(
        self,
        file_paths: List[str],
        metadata_template: Optional[DocumentMetadata] = None
    ) -> IngestionJob:
        """Ingest a list of files."""
        return await self.batch_service.process_files(file_paths, metadata_template)