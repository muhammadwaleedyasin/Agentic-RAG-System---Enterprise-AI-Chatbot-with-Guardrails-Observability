"""
Document processing service for text extraction and chunking.
"""
import asyncio
import uuid
from typing import List, Dict, Any, Optional
import logging
from pathlib import Path
import mimetypes

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import markdown

from ..config.settings import settings
from ..models.documents import Document, DocumentChunk, DocumentType, DocumentMetadata

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing and chunking documents."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.max_file_size = settings.max_file_size
        self.allowed_file_types = settings.allowed_file_types
    
    async def process_file(self, file_path: str, metadata: DocumentMetadata = None) -> Document:
        """
        Process a file and extract text content.
        
        Args:
            file_path: Path to the file to process
            metadata: Optional document metadata
            
        Returns:
            Processed document with extracted content
        """
        try:
            file_path = Path(file_path)
            
            # Validate file
            await self._validate_file(file_path)
            
            # Determine document type
            doc_type = self._get_document_type(file_path)
            
            # Extract content
            content = await self._extract_content(file_path, doc_type)
            
            # Create document metadata
            if metadata is None:
                metadata = DocumentMetadata()
            
            # Update metadata with file info
            metadata.file_size = file_path.stat().st_size
            if not metadata.title:
                metadata.title = file_path.stem
            
            # Create document
            document = Document(
                document_id=str(uuid.uuid4()),
                filename=file_path.name,
                content=content,
                document_type=doc_type,
                status="completed",
                metadata=metadata
            )
            
            logger.info(f"Processed document: {file_path.name} ({len(content)} characters)")
            return document
            
        except Exception as e:
            logger.error(f"Failed to process file {file_path}: {str(e)}")
            raise
    
    async def process_text(self, text: str, filename: str = None, metadata: DocumentMetadata = None) -> Document:
        """
        Process raw text content.
        
        Args:
            text: Text content to process
            filename: Optional filename
            metadata: Optional document metadata
            
        Returns:
            Processed document
        """
        try:
            # Create document metadata
            if metadata is None:
                metadata = DocumentMetadata()
            
            if not metadata.title and filename:
                metadata.title = filename
            
            # Create document
            document = Document(
                document_id=str(uuid.uuid4()),
                filename=filename or "text_document.txt",
                content=text,
                document_type=DocumentType.TEXT,
                status="completed",
                metadata=metadata
            )
            
            logger.info(f"Processed text document: {len(text)} characters")
            return document
            
        except Exception as e:
            logger.error(f"Failed to process text: {str(e)}")
            raise
    
    async def chunk_document(self, document: Document) -> List[DocumentChunk]:
        """
        Split document content into chunks.
        
        Args:
            document: Document to chunk
            
        Returns:
            List of document chunks
        """
        try:
            if not document.content:
                return []
            
            chunks = []
            content = document.content
            chunk_index = 0
            
            # Split content into chunks
            for i in range(0, len(content), self.chunk_size - self.chunk_overlap):
                # Calculate chunk boundaries
                start_char = i
                end_char = min(i + self.chunk_size, len(content))
                
                # Extract chunk content
                chunk_content = content[start_char:end_char]
                
                # Skip very small chunks
                if len(chunk_content.strip()) < 10:
                    continue
                
                # Create chunk
                chunk = DocumentChunk(
                    chunk_id=f"{document.document_id}_{chunk_index}",
                    document_id=document.document_id,
                    content=chunk_content,
                    chunk_index=chunk_index,
                    start_char=start_char,
                    end_char=end_char,
                    metadata={
                        "document_title": document.metadata.title,
                        "document_type": document.document_type.value,
                        "filename": document.filename
                    }
                )
                
                chunks.append(chunk)
                chunk_index += 1
                
                # Stop if we've reached the end
                if end_char >= len(content):
                    break
            
            logger.info(f"Created {len(chunks)} chunks for document {document.document_id}")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to chunk document: {str(e)}")
            raise
    
    async def _validate_file(self, file_path: Path):
        """Validate file size and type."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise ValueError(f"File too large: {file_size} bytes (max: {self.max_file_size})")
        
        # Check file type
        file_extension = file_path.suffix.lower()
        if file_extension not in self.allowed_file_types:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _get_document_type(self, file_path: Path) -> DocumentType:
        """Determine document type from file extension."""
        extension = file_path.suffix.lower()
        
        type_mapping = {
            ".txt": DocumentType.TEXT,
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".md": DocumentType.MARKDOWN,
            ".html": DocumentType.HTML,
        }
        
        return type_mapping.get(extension, DocumentType.TEXT)
    
    async def _extract_content(self, file_path: Path, doc_type: DocumentType) -> str:
        """Extract text content from file based on document type."""
        loop = asyncio.get_event_loop()
        
        if doc_type == DocumentType.TEXT:
            return await loop.run_in_executor(None, self._extract_text, file_path)
        elif doc_type == DocumentType.PDF:
            return await loop.run_in_executor(None, self._extract_pdf, file_path)
        elif doc_type == DocumentType.DOCX:
            return await loop.run_in_executor(None, self._extract_docx, file_path)
        elif doc_type == DocumentType.MARKDOWN:
            return await loop.run_in_executor(None, self._extract_markdown, file_path)
        elif doc_type == DocumentType.HTML:
            return await loop.run_in_executor(None, self._extract_html, file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract content from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {str(e)}")
            raise ValueError(f"Failed to process PDF file: {str(e)}")
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {str(e)}")
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    def _extract_markdown(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to plain text
            # Remove markdown formatting for better chunking
            text = markdown.markdown(md_content)
            # Remove HTML tags (simple approach)
            import re
            text = re.sub('<[^<]+?>', '', text)
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract Markdown content: {str(e)}")
            raise ValueError(f"Failed to process Markdown file: {str(e)}")
    
    def _extract_html(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
        except ImportError:
            raise ValueError("BeautifulSoup4 is required for HTML processing")
        except Exception as e:
            logger.error(f"Failed to extract HTML content: {str(e)}")
            raise ValueError(f"Failed to process HTML file: {str(e)}")
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported file types."""
        return self.allowed_file_types.copy()
    
    async def health_check(self) -> bool:
        """Check if the document processor is healthy."""
        try:
            # Test basic text processing
            test_doc = await self.process_text("Test content", "test.txt")
            chunks = await self.chunk_document(test_doc)
            return len(chunks) > 0
        except Exception:
            return False